#!/usr/bin/env python3
"""
Compile all analysis sections into a CEO-ready Word document.
Reads markdown files and agent JSONL outputs, converts to formatted .docx.
"""

import json
import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = "/Users/matt/Documents/monsterbass"
OUTPUT_DIR = f"{BASE}/output"

# ── Helpers ──────────────────────────────────────────────────────────

def extract_analysis_from_jsonl(filepath):
    """Extract the longest text block from an agent's JSONL output file."""
    longest = ""
    try:
        with open(filepath, "r") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    obj = json.loads(line)
                except json.JSONDecodeError:
                    continue
                msg = obj.get("message", {})
                content = msg.get("content", [])
                if isinstance(content, list):
                    for block in content:
                        if isinstance(block, dict) and block.get("type") == "text":
                            text = block.get("text", "")
                            if len(text) > len(longest):
                                longest = text
    except FileNotFoundError:
        pass
    return longest


def read_md(filepath):
    """Read a markdown file."""
    with open(filepath, "r") as f:
        return f.read()


def set_cell_shading(cell, color):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_text(paragraph, text):
    """Add text with basic bold/italic markdown formatting to a paragraph."""
    # Process **bold** and *italic* patterns
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def parse_table(lines):
    """Parse markdown table lines into list of lists."""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and not re.match(r'^\|[\s\-:|]+\|$', line):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows.append(cells)
    return rows


def add_table_to_doc(doc, rows, first_row_header=True):
    """Add a formatted table to the document."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    # Pad rows
    for r in rows:
        while len(r) < n_cols:
            r.append("")

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.style = doc.styles['Normal']
            # Strip markdown bold from cell text
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
            clean = re.sub(r'\*(.*?)\*', r'\1', clean)
            run = p.add_run(clean)
            run.font.size = Pt(8)
            if i == 0 and first_row_header:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_shading(cell, "2F5496")
            elif i % 2 == 0 and i > 0:
                set_cell_shading(cell, "D6E4F0")

    doc.add_paragraph()  # spacing


def add_section_from_markdown(doc, md_text, skip_title=False):
    """Parse markdown text and add formatted content to the Word document."""
    lines = md_text.split('\n')
    i = 0
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip the first H1 title if requested (we add our own)
        if skip_title and stripped.startswith('# ') and not stripped.startswith('## '):
            i += 1
            # Skip metadata lines right after title
            while i < len(lines) and (lines[i].strip().startswith('**') or lines[i].strip() == '---' or lines[i].strip() == ''):
                i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            i += 1
            continue

        # Empty line
        if not stripped:
            # Flush table
            if in_table and table_lines:
                rows = parse_table(table_lines)
                add_table_to_doc(doc, rows)
                table_lines = []
                in_table = False
            i += 1
            continue

        # Table line
        if stripped.startswith('|'):
            in_table = True
            table_lines.append(stripped)
            i += 1
            continue
        else:
            # Flush any pending table
            if in_table and table_lines:
                rows = parse_table(table_lines)
                add_table_to_doc(doc, rows)
                table_lines = []
                in_table = False

        # Headers
        if stripped.startswith('### '):
            p = doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        elif stripped.startswith('# '):
            p = doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            i += 1
            continue

        # Code block
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # skip closing ```
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.size = Pt(8)
            run.font.name = 'Courier New'
            pf = p.paragraph_format
            pf.left_indent = Cm(1)
            continue

        # Checkbox items (- [ ] ...)
        if stripped.startswith('- [ ]') or stripped.startswith('- [x]'):
            checked = stripped.startswith('- [x]')
            text = stripped[6:].strip()
            marker = "[x]" if checked else "[ ]"
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f"{marker} ")
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            add_formatted_text(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1

    # Flush final table
    if in_table and table_lines:
        rows = parse_table(table_lines)
        add_table_to_doc(doc, rows)


# ── Main ─────────────────────────────────────────────────────────────

def main():
    # Collect content from all sources
    print("Reading analysis files...")

    books_md = read_md(f"{OUTPUT_DIR}/Books_Quality_Assessment.md")
    customer_md = read_md(f"{OUTPUT_DIR}/CEO_Customer_Revenue_Deep_Dive.md")

    # Extract from agent JSONL outputs
    tasks_dir = "/private/tmp/claude-501/-Users-matt-Documents-monsterbass/tasks"
    vendor_text = extract_analysis_from_jsonl(f"{tasks_dir}/a91af03.output")
    strategic_text = extract_analysis_from_jsonl(f"{tasks_dir}/ae341ec.output")

    # Save extracted analyses to md files for reference
    if vendor_text:
        with open(f"{OUTPUT_DIR}/Vendor_Expense_Analysis.md", "w") as f:
            f.write(vendor_text)
        print(f"  Vendor analysis: {len(vendor_text):,} chars")
    else:
        print("  WARNING: Could not extract vendor analysis from agent output")
        vendor_text = "Vendor expense analysis not available - see agent output."

    if strategic_text:
        with open(f"{OUTPUT_DIR}/Strategic_CFO_Roadshow.md", "w") as f:
            f.write(strategic_text)
        print(f"  Strategic analysis: {len(strategic_text):,} chars")
    else:
        print("  WARNING: Could not extract strategic analysis from agent output")
        strategic_text = "Strategic analysis not available - see agent output."

    print(f"  Books assessment: {len(books_md):,} chars")
    print(f"  Customer deep dive: {len(customer_md):,} chars")

    # Create Word document
    print("\nBuilding Word document...")
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # ── Style adjustments ──
    style = doc.styles['Normal']
    style.font.size = Pt(10)
    style.font.name = 'Calibri'
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Calibri'
        if level == 1:
            h_style.font.size = Pt(18)
            h_style.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        elif level == 2:
            h_style.font.size = Pt(14)
            h_style.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        elif level == 3:
            h_style.font.size = Pt(11)
            h_style.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    # ── Cover Page ──
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MONSTERBASS")
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Outdoor Playground Inc")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("Financial & Operational Audit")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    doc.add_paragraph()

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("CEO Briefing Document")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    for _ in range(4):
        doc.add_paragraph()

    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = details.add_run("Audit Period: January 2024 -- December 2025\n")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
    run = details.add_run("Data Sources: QuickBooks Online, Shopify Admin API, Shopify Capital Records\n")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
    run = details.add_run("Report Date: February 5, 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    # ── Page break ──
    doc.add_page_break()

    # ── Table of Contents (manual) ──
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        ("1.", "Executive Summary", 2),
        ("2.", "Quality of the Books", 4),
        ("3.", "Vendor Expense Analysis", 10),
        ("4.", "Customer & Revenue Deep Dive", 16),
        ("5.", "Strategic CFO Insights & Roadshow Preparation", 24),
    ]
    for num, title_text, _ in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}  {title_text}")
        run.font.size = Pt(12)
        run.bold = True

    doc.add_page_break()

    # ── Section 1: Executive Summary ──
    doc.add_heading("1. Executive Summary", level=1)

    exec_summary = """This document presents the findings of a comprehensive financial and operational audit of MonsterBass (Outdoor Playground Inc), covering the period January 2024 through December 2025. The audit integrates data from QuickBooks Online (general ledger, balance sheet, P&L) with granular Shopify transaction data (101,126 orders, 258,262 customer records, 871 payouts) and Shopify Capital loan records."""
    p = doc.add_paragraph()
    add_formatted_text(p, exec_summary)

    doc.add_heading("Key Findings at a Glance", level=2)

    summary_items = [
        ("Revenue Recognition", "The bookkeeper records net Shopify payouts as revenue, understating actual gross revenue by approximately 28% ($700K+ on Shopify alone). This is the single largest issue and cascades into every other financial metric."),
        ("COGS & Inventory", "Only $43.6K in COGS recorded over 24 months for a $3M+ revenue business. Inventory on the balance sheet ($1.294M) has never been relieved -- it grows every month. Estimated overstatement: $500K-$1M."),
        ("Revenue Trend", "2025 full-year Shopify revenue declined 11.4% YoY ($1.74M vs $1.96M in 2024). Q1 2025 was worst (-23% to -30% YoY), but Nov-Dec 2025 turned positive (+14.8%, +5.9%) driven by TikTok Shop explosion."),
        ("Marketing Spend", "Total marketing + creative ecosystem: $677K in 2025, representing 39.8% of revenue. Industry norm is 15-20%. Six to seven separate content contractors create significant redundancy."),
        ("Customer Economics", "Average LTV: $192, Median: $82. 58.2% repeat purchase rate (strong). Top 10% of customers drive 44.9% of revenue. Subscription boxes are 74.4% of revenue."),
        ("Cash Position", "Cash on hand: $100K. Monthly burn: $10-30K negative. Shopify Capital ($140K loan, $67K remaining) deducts 13% of daily sales before payout. Estimated runway: 2-4 months at current trajectory."),
        ("Accounts Payable", "$981K in AP, with 93.7% over 91 days past due. This is the most acute distress signal."),
        ("Refunds", "2.42% of revenue -- healthy for D2C subscription. Improved significantly in mid-2025 (below 1% in several months)."),
    ]

    for title_text, detail in summary_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(detail)
        run2.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    add_formatted_text(p, "**Bottom line:** The books are functional for cash tracking but not GAAP-compliant and not suitable for investor or buyer diligence in their current state. The underlying business has genuine strengths (strong repeat rate, subscription model, passionate niche market) but is under financial stress. The path forward requires 12-18 weeks of financial remediation, immediate cost rationalization (potential $315K-$519K in annual savings identified), and rebuilding the customer acquisition engine.")

    doc.add_page_break()

    # ── Section 2: Quality of the Books ──
    doc.add_heading("2. Quality of the Books", level=1)
    add_section_from_markdown(doc, books_md, skip_title=True)

    doc.add_page_break()

    # ── Section 3: Vendor Expense Analysis ──
    doc.add_heading("3. Vendor Expense Analysis (2025)", level=1)
    add_section_from_markdown(doc, vendor_text, skip_title=True)

    doc.add_page_break()

    # ── Section 4: Customer & Revenue Deep Dive ──
    doc.add_heading("4. Customer & Revenue Deep Dive", level=1)
    add_section_from_markdown(doc, customer_md, skip_title=True)

    doc.add_page_break()

    # ── Section 5: Strategic CFO & Roadshow ──
    doc.add_heading("5. Strategic CFO Insights & Roadshow Preparation", level=1)
    add_section_from_markdown(doc, strategic_text, skip_title=True)

    # ── Save ──
    output_path = f"{BASE}/output/MonsterBass_CEO_Audit_Briefing.docx"
    doc.save(output_path)
    print(f"\nDocument saved to: {output_path}")
    print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")


if __name__ == "__main__":
    main()
